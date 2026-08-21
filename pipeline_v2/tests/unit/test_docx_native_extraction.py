from __future__ import annotations

import base64
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from rag_pipeline.contracts import (
    ElementType,
    EngineName,
    FileFormat,
    RouteDecision,
    RouteType,
)
from rag_pipeline.extractors import (
    DocxNativeExtractionError,
    DocxNativeExtractor,
    ExtractionContractError,
    ExtractionRequest,
    provenance_coverage,
    sha256_file,
)


_ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwC"
    "AAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def docx_route() -> RouteDecision:
    return RouteDecision(
        route_id="route-docx-native",
        catalog_id="catalog-docx",
        version_id="family-docx-v001",
        file_format=FileFormat.DOCX,
        route_type=RouteType.DOCX_NATIVE,
        primary_engine=EngineName.DOCX_NATIVE,
        fallback_engines=[],
        reason_code="docx_native_primary",
    )


def request_for(path: Path) -> ExtractionRequest:
    return ExtractionRequest(
        source_path=path,
        catalog_id="catalog-docx",
        family_id="family-docx",
        version_id="family-docx-v001",
        route=docx_route(),
        run_id="docx-unit-test",
    )


def build_rich_docx(path: Path) -> None:
    document = Document()
    document.sections[0].header.paragraphs[0].text = "Petrolimex Aviation"
    document.sections[0].footer.paragraphs[0].text = "Tài liệu nội bộ"
    document.add_paragraph("Quy trình quản lý vật tư", style="Title")
    document.add_heading("Phạm vi", level=1)
    document.add_paragraph("Áp dụng cho toàn bộ đơn vị khai thác.")
    document.add_paragraph("Lập đề nghị", style="List Number")
    document.add_paragraph("Phê duyệt đề nghị", style="List Number")
    table = document.add_table(rows=2, cols=2)
    merged = table.cell(0, 0).merge(table.cell(0, 1))
    merged.text = "Biểu mẫu"
    table.cell(1, 0).text = "Mã"
    table.cell(1, 1).text = "BM.QT.TCKT.03"
    document.add_paragraph("Bảng 1. Danh mục biểu mẫu", style="Caption")
    document.add_section()
    document.add_heading("Lưu hồ sơ", level=1)
    document.add_paragraph("Hồ sơ được lưu theo quy định.")
    document.add_picture(BytesIO(_ONE_PIXEL_PNG))
    document.save(path)


def test_docx_native_returns_unified_document_with_provenance(
    tmp_path: Path,
) -> None:
    source = tmp_path / "procedure.docx"
    build_rich_docx(source)
    source_before = source.read_bytes()

    result = DocxNativeExtractor().extract(request_for(source))
    extracted = result.document

    assert source.read_bytes() == source_before
    assert extracted.catalog_id == "catalog-docx"
    assert extracted.family_id == "family-docx"
    assert extracted.version_id == "family-docx-v001"
    assert extracted.source_sha256 == sha256_file(source)
    assert extracted.page_count is None
    assert extracted.engine_chain == [EngineName.DOCX_NATIVE]
    assert provenance_coverage(extracted) == 1.0
    assert result.metrics["external_calls"] == 0
    assert result.metrics["ocr_invocations"] == 0
    assert result.metrics["section_count"] == 2

    types = [element.element_type for element in extracted.elements]
    assert ElementType.TITLE in types
    assert ElementType.HEADING in types
    assert ElementType.PARAGRAPH in types
    assert ElementType.LIST in types
    assert ElementType.TABLE in types
    assert ElementType.CAPTION in types
    assert ElementType.IMAGE in types
    assert ElementType.HEADER in types
    assert ElementType.FOOTER in types

    list_element = next(
        element
        for element in extracted.elements
        if element.element_type == ElementType.LIST
    )
    assert "1. Lập đề nghị" in list_element.content
    assert "2. Phê duyệt đề nghị" in list_element.content
    assert list_element.metadata["item_count"] == 2

    table_element = next(
        element
        for element in extracted.elements
        if element.element_type == ElementType.TABLE
    )
    assert "| Biểu mẫu |  |" in table_element.content
    assert "| Mã | BM.QT.TCKT.03 |" in table_element.content
    assert table_element.metadata["gfm_table"] is True
    assert table_element.metadata["merged_slot_count"] >= 1

    assert len(extracted.assets) == 1
    assert extracted.assets[0].references
    assert extracted.assets[0].retain_raw is True
    assert extracted.assets[0].indexable is False


def test_docx_native_ids_are_stable_for_same_input(tmp_path: Path) -> None:
    source = tmp_path / "stable.docx"
    build_rich_docx(source)
    extractor = DocxNativeExtractor()

    first = extractor.extract(request_for(source)).document
    second = extractor.extract(request_for(source)).document

    assert [element.element_id for element in first.elements] == [
        element.element_id for element in second.elements
    ]
    assert [
        element.source_anchor.anchor_id
        for element in first.elements
    ] == [
        element.source_anchor.anchor_id
        for element in second.elements
    ]
    assert [asset.asset_id for asset in first.assets] == [
        asset.asset_id for asset in second.assets
    ]


def test_docx_native_rejects_invalid_package(tmp_path: Path) -> None:
    source = tmp_path / "broken.docx"
    source.write_bytes(b"not-a-docx-package")

    with pytest.raises(DocxNativeExtractionError):
        DocxNativeExtractor().extract(request_for(source))


def test_docx_extractor_rejects_pdf_route(tmp_path: Path) -> None:
    source = tmp_path / "source.pdf"
    source.write_bytes(b"%PDF-1.7\n")
    route = RouteDecision(
        route_id="route-pdf",
        catalog_id="catalog-docx",
        version_id="family-docx-v001",
        file_format=FileFormat.PDF,
        route_type=RouteType.PDF_OPENDATALOADER,
        primary_engine=EngineName.OPENDATALOADER_LOCAL,
        fallback_engines=[EngineName.OPENDATALOADER_HYBRID],
        reason_code="pdf_primary",
    )
    request = ExtractionRequest(
        source_path=source,
        catalog_id="catalog-docx",
        family_id="family-docx",
        version_id="family-docx-v001",
        route=route,
    )

    with pytest.raises(ExtractionContractError):
        DocxNativeExtractor().extract(request)
