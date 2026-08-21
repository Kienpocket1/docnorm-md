from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document

from rag_pipeline.contracts import (
    ElementType,
    EngineName,
    ExtractedDocument,
    FileFormat,
    ResumeFrom,
    RouteDecision,
    RouteType,
    SourceAnchor,
    SourceElement,
    StageStatus,
)
from rag_pipeline.extractors import ExtractionRequest, sha256_file
from rag_pipeline.stages.extract import (
    ExtractStage,
    ExtractStageConfigurationError,
    UnsupportedExtractionRouteError,
)


def route(
    *,
    file_format: FileFormat,
    route_type: RouteType,
    engine: EngineName,
) -> RouteDecision:
    return RouteDecision(
        route_id=f"route-{route_type.value.lower()}",
        catalog_id="catalog-001",
        version_id="family-001-v001",
        file_format=file_format,
        route_type=route_type,
        primary_engine=engine,
        fallback_engines=(
            [EngineName.OPENDATALOADER_HYBRID]
            if route_type == RouteType.PDF_OPENDATALOADER
            else []
        ),
        reason_code="unit_test_route",
    )


def request(path: Path, decision: RouteDecision) -> ExtractionRequest:
    return ExtractionRequest(
        source_path=path,
        catalog_id="catalog-001",
        family_id="family-001",
        version_id="family-001-v001",
        route=decision,
        run_id="extract-stage-test",
    )


def test_docx_route_returns_stage_result_without_quality(
    tmp_path: Path,
) -> None:
    source = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("Quy trình", level=1)
    document.add_paragraph("Nội dung kiểm thử DOCX native.")
    document.save(source)
    outcome = ExtractStage().run(
        request(
            source,
            route(
                file_format=FileFormat.DOCX,
                route_type=RouteType.DOCX_NATIVE,
                engine=EngineName.DOCX_NATIVE,
            ),
        )
    )

    assert isinstance(outcome.extraction.document, ExtractedDocument)
    assert outcome.stage_result.stage == ResumeFrom.EXTRACT
    assert outcome.stage_result.status == StageStatus.PASS
    assert outcome.stage_result.metrics["quality_gate_invocations"] == 0
    assert outcome.stage_result.metrics["fallback_invocations"] == 0
    assert outcome.stage_result.message is not None
    assert "EXTRACTION_QUALITY" in outcome.stage_result.message


class FakePdfExtractor:
    def __init__(self, source: Path, raw_directory: Path) -> None:
        self.source = source
        self.raw_directory = raw_directory
        self.calls = 0

    def extract(self, input_pdf, *, run_id=None):
        self.calls += 1
        assert Path(input_pdf).resolve() == self.source.resolve()
        markdown = self.raw_directory / "document.md"
        structured = self.raw_directory / "document.json"
        markdown.write_text("# PDF\n", encoding="utf-8")
        structured.write_text("{}", encoding="utf-8")
        result_manifest = self.raw_directory / "worker_result.json"
        result_manifest.write_text("{}", encoding="utf-8")
        stdout = self.raw_directory / "stdout.log"
        stderr = self.raw_directory / "stderr.log"
        stdout.write_text("ok", encoding="utf-8")
        stderr.write_text("", encoding="utf-8")
        record = lambda path: SimpleNamespace(path=path)
        manifest = SimpleNamespace(
            worker_version="fake-worker",
            engine_version="2.5.0",
            elapsed_seconds=0.01,
            output=SimpleNamespace(
                markdown=record(markdown),
                structured_json=record(structured),
                images=[],
            ),
        )
        return SimpleNamespace(
            manifest=manifest,
            result_manifest_path=result_manifest,
            stdout_log_path=stdout,
            stderr_log_path=stderr,
        )


class FakePdfMapper:
    def __init__(self, source: Path) -> None:
        self.source = source
        self.calls = 0

    def map_run(self, run, identity):
        self.calls += 1
        source_sha256 = sha256_file(self.source)
        anchor = SourceAnchor(
            anchor_id="anchor-pdf-1",
            source_sha256=source_sha256,
            page_number=1,
            block_index=0,
            bbox=None,
            section_path=["PDF"],
        )
        document = ExtractedDocument(
            catalog_id=identity.catalog_id,
            family_id=identity.family_id,
            version_id=identity.version_id,
            source_sha256=source_sha256,
            page_count=1,
            engine_chain=[EngineName.OPENDATALOADER_LOCAL],
            elements=[
                SourceElement(
                    element_id="element-pdf-1",
                    element_type=ElementType.PARAGRAPH,
                    content="PDF body",
                    source_anchor=anchor,
                    engine=EngineName.OPENDATALOADER_LOCAL,
                    metadata={"indexable": True},
                )
            ],
        )
        return SimpleNamespace(
            document=document,
            mapping_issues=(),
            metrics={"mapped_element_count": 1},
        )


def test_pdf_route_uses_local_worker_and_shared_contract(
    tmp_path: Path,
) -> None:
    source = tmp_path / "sample.pdf"
    source.write_bytes(b"%PDF-1.7\n% stage 5 unit test\n")
    raw_directory = tmp_path / "raw"
    raw_directory.mkdir()
    extractor = FakePdfExtractor(source, raw_directory)
    mapper = FakePdfMapper(source)
    outcome = ExtractStage(
        pdf_extractor=extractor,
        pdf_mapper=mapper,
    ).run(
        request(
            source,
            route(
                file_format=FileFormat.PDF,
                route_type=RouteType.PDF_OPENDATALOADER,
                engine=EngineName.OPENDATALOADER_LOCAL,
            ),
        )
    )

    assert extractor.calls == 1
    assert mapper.calls == 1
    assert isinstance(outcome.extraction.document, ExtractedDocument)
    assert outcome.extraction.document.page_count == 1
    assert outcome.stage_result.status == StageStatus.PASS
    assert outcome.stage_result.metrics["document_format"] == "PDF"
    assert outcome.stage_result.metrics["provenance_coverage"] == 1.0
    assert outcome.stage_result.metrics["mapping_issue_count"] == 0


def test_pdf_route_requires_both_adapter_and_mapper(tmp_path: Path) -> None:
    source = tmp_path / "sample.pdf"
    source.write_bytes(b"%PDF-1.7\n")

    with pytest.raises(ExtractStageConfigurationError):
        ExtractStage().run(
            request(
                source,
                route(
                    file_format=FileFormat.PDF,
                    route_type=RouteType.PDF_OPENDATALOADER,
                    engine=EngineName.OPENDATALOADER_LOCAL,
                ),
            )
        )


def test_stage_rejects_route_owned_by_another_stage(tmp_path: Path) -> None:
    source = tmp_path / "legacy.doc"
    source.write_bytes(b"legacy-doc-placeholder")
    decision = route(
        file_format=FileFormat.DOC,
        route_type=RouteType.DOC_CONVERT,
        engine=EngineName.LIBREOFFICE,
    )

    with pytest.raises(UnsupportedExtractionRouteError):
        ExtractStage().run(request(source, decision))
