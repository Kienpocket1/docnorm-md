from __future__ import annotations

from pathlib import Path

import pymupdf
import pytest
from docx import Document

from docnorm_web.models import RepairProfile
from docnorm_web.normalizer import (
    ScanRequiresOcrError,
    analyze_pdf_text_layer,
    clean_markdown,
    extract_docx_basic,
    extract_pdf_basic,
)
from docnorm_web.security import InvalidDocumentError, validate_document
from docnorm_web.structural_profiles import (
    FLOWCHART_NOTE_1,
    FLOWCHART_NOTE_2,
    StructuralProfileError,
    apply_structural_profile,
    count_gfm_tables,
    validate_petrolimex_materials_markdown,
)

AUDITED_MATERIALS_MARKDOWN = """# QUY TRÌNH QUẢN LÝ VẬT TƯ

## (QT.TCKT.01)

a. x

# QUY TRÌNH QUẢN LÝ VẬT TƯ

i. I. MỤC ĐÍCH: Quản lý vật tư.
ii. II. PHẠM VI ÁP DỤNG:
- - Toàn Công ty Cổ phần Nhiên liệu bay Petrolimex.
iii. III. ĐỊNH NGHĨA VÀ CÁC TỪ VIẾT TẮT:
  vụ cho sản xuất kinh doanh.
  - Kho vật tư: Kho vật tư của Công ty tại Chi nhánh.
  - Công ty: Công ty Cổ phần Nhiên liệu bay Petrolimex.
  - P.KT-TH: Phòng Kế toán - Tổng hợp.
  - Vật tư: Nguyên liệu, vật liệu, máy móc, phụ tùng thay thế phục
iv. IV. Qui định chung:
- - Vật tư phải được quản lý.

## V. LƯU ĐỒ

1. Nhập vật tư

Tiếp nhận vật tư, hồ sơ

Không đạt

Kiểm tra, xử lý

Đạt

Lưu kho

2. Xuất cấp vật tư

Thông tin đề xuất

Không bàn giao

Bàn giao

i. VI. DIỄN GIẢI LƯU ĐỒ

1. 1. Nhập vật tư

| TT | Nội dung | Diễn giải | Hồ sơ |
| --- | --- | --- | --- |
| 1 | Tiếp nhận | Nội dung 1 | BM.QT.TCKT.01 |
| 2 | Kiểm tra | Nội dung 2 | BM.QT.TCKT.02 |

| 3 | Bàn giao | Nội dung 3 | 1. BM.QT.TCKT.01,<br>2. BM.QT.TCKT.02 |
| --- | --- | --- | --- |
| 4 | Theo dõi | Nội dung 4 | BM.QT.TCKT.05 |
| 5 | Báo cáo | Nội dung 5 | BM.QT.TCKT.05 |

## 2. Xuất cấp vật tư

| TT | Nội dung | Diễn giải | Hồ sơ |
| --- | --- | --- | --- |
| 1 | Đề xuất | Nội dung 1 | BM.QT.TCKT.03 |
| 2 | Xử lý | Nội dung 2 |  |
| 3 | Bàn giao | Nội dung 3 | BM.QT.TCKT.04 |

| 4 | Theo dõi | Nội dung 4 | BM.QT.TCKT.05 |
| --- | --- | --- | --- |
| 5 | Báo cáo | Nội dung 5 | BM.QT.TCKT.05 |

## 3. Luân chuyển chứng từ và kiểm tra báo cáo

| TT | Trách nhiệm | Nội dung hoạt động | Hồ sơ |
| --- | --- | --- | --- |
| 1 | P.KT-TH | Nội dung đầu.<br>A. − Nội dung hai.<br>B. − Nội dung ba. | BM.QT.TCKT.05 |

BM.QT.TCKT.06

## Đại diện bên giao Đại diện bên nhận

## ………., ngày tháng năm …… DUYỆT P.KT-TH QUẢN LÝ KHO NGƯỜI NHẬN

PETROLIMEX AVIATION

QUY TRÌNH QUẢN LÝ VẬT TƯ
"""


def test_clean_markdown_removes_double_bullet_and_normalizes_blank_lines() -> None:
    assert clean_markdown("- - Một\n\n\n\nHai\r\n") == "- Một\n\nHai\n"


def test_docx_basic_preserves_heading_paragraph_list_and_table(tmp_path: Path) -> None:
    source = tmp_path / "demo.docx"
    document = Document()
    document.add_heading("Quy trình mẫu", level=1)
    document.add_paragraph("Nội dung kiểm thử UTF-8.")
    document.add_paragraph("Mục thứ nhất", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Cột A"
    table.cell(0, 1).text = "Cột B"
    table.cell(1, 0).text = "Một"
    table.cell(1, 1).text = "Hai"
    document.save(source)
    validate_document(source, 10 * 1024 * 1024)
    result = extract_docx_basic(source)
    assert "# Quy trình mẫu" in result.markdown
    assert "| Cột A | Cột B |" in result.markdown
    assert result.metrics.tables == 1
    assert result.metrics.broken_characters == 0
    assert result.backend == "BASIC_FALLBACK"


def test_pdf_basic_extracts_text_and_page_provenance(tmp_path: Path) -> None:
    source = tmp_path / "demo.pdf"
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 80), "DOCNORM MD DEMO", fontsize=18)
    page.insert_text(
        (72, 120), "This PDF contains enough selectable text for extraction."
    )
    document.save(source)
    document.close()
    validate_document(source, 10 * 1024 * 1024)
    result = extract_pdf_basic(source)
    assert "<!-- page: 1 -->" in result.markdown
    assert "DOCNORM MD DEMO" in result.markdown
    assert result.metrics.pages == 1
    assert result.metrics.provenance_coverage == 1.0


def test_textless_pdf_fails_closed_for_ocr(tmp_path: Path) -> None:
    source = tmp_path / "scan.pdf"
    document = pymupdf.open()
    document.new_page()
    document.save(source)
    document.close()
    with pytest.raises(ScanRequiresOcrError):
        extract_pdf_basic(source)


def test_image_scan_is_routed_to_ocr_but_text_pdf_is_not(tmp_path: Path) -> None:
    scan = tmp_path / "image-scan.pdf"
    document = pymupdf.open()
    page = document.new_page()
    pixmap = pymupdf.Pixmap(pymupdf.csRGB, pymupdf.IRect(0, 0, 200, 280), False)
    pixmap.clear_with(245)
    page.insert_image(page.rect, pixmap=pixmap)
    document.save(scan)
    document.close()
    scan_analysis = analyze_pdf_text_layer(scan)
    assert scan_analysis.page_count == 1
    assert scan_analysis.image_dominant_pages == 1
    assert scan_analysis.requires_ocr is True

    text_pdf = tmp_path / "text.pdf"
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text(
        (72, 90),
        "Selectable document text with enough characters for normal extraction.",
    )
    document.save(text_pdf)
    document.close()
    text_analysis = analyze_pdf_text_layer(text_pdf)
    assert text_analysis.text_pages == 1
    assert text_analysis.requires_ocr is False


def test_magic_mismatch_is_rejected(tmp_path: Path) -> None:
    source = tmp_path / "fake.pdf"
    source.write_text("not a PDF", encoding="utf-8")
    with pytest.raises(InvalidDocumentError):
        validate_document(source, 1024)


def test_materials_profile_repairs_every_audited_issue_and_is_idempotent() -> None:
    result = apply_structural_profile(
        AUDITED_MATERIALS_MARKDOWN,
        RepairProfile.AUTO,
    )
    assert result.applied == RepairProfile.PETROLIMEX_MATERIALS
    assert result.fingerprint_score == 1.0
    assert len(result.resolved_issue_codes) == 9
    assert "a. x" not in result.markdown
    assert "PETROLIMEX AVIATION" not in result.markdown
    assert "i. I." not in result.markdown
    assert "- -" not in result.markdown
    assert "<br>A. −" not in result.markdown
    assert "1. BM.QT.TCKT.01" not in result.markdown
    assert "## Đại diện bên giao" not in result.markdown
    assert result.markdown.count("QUY TRÌNH QUẢN LÝ VẬT TƯ") == 1
    assert result.markdown.count("| TT | Nội dung | Diễn giải | Hồ sơ |") == 2
    assert FLOWCHART_NOTE_1 in result.markdown
    assert FLOWCHART_NOTE_2 in result.markdown
    assert count_gfm_tables(result.markdown) == 5
    validate_petrolimex_materials_markdown(result.markdown)

    second = apply_structural_profile(
        result.markdown,
        RepairProfile.PETROLIMEX_MATERIALS,
    )
    assert second.markdown == result.markdown


def test_auto_profile_never_changes_unrelated_document() -> None:
    markdown = "# Quy trình nghỉ phép\n\nNội dung độc lập.\n"
    result = apply_structural_profile(markdown, RepairProfile.AUTO)
    assert result.applied == RepairProfile.GENERIC
    assert result.markdown == markdown
    assert result.resolved_issue_codes == ()


def test_forced_materials_profile_fails_closed_on_mismatch() -> None:
    with pytest.raises(StructuralProfileError):
        apply_structural_profile(
            "# Tài liệu khác\n",
            RepairProfile.PETROLIMEX_MATERIALS,
        )
