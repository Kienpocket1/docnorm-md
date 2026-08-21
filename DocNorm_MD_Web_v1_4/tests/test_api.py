from __future__ import annotations

import io
import time
import zipfile

from docx import Document
from fastapi.testclient import TestClient

from docnorm_web.main import app

client = TestClient(app)


def _docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Tài liệu demo", level=1)
    document.add_paragraph("Nội dung dùng để kiểm thử API DocNorm MD.")
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_health_is_local_product() -> None:
    page = client.get("/")
    assert page.status_code == 200
    assert "DocNorm MD" in page.text
    assert "Kéo thả PDF hoặc DOCX" in page.text
    assert "Content-Security-Policy" in page.headers
    response = client.get("/api/health")
    assert response.status_code == 200
    payload = response.json()
    assert payload["product"] == "DocNorm MD"
    assert payload["supported_formats"] == ["pdf", "docx"]


def test_docx_job_completes_and_downloads_closed_bundle() -> None:
    response = client.post(
        "/api/jobs",
        data={"mode": "fast"},
        files={
            "file": (
                "demo.docx",
                _docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 202
    job_id = response.json()["job_id"]
    payload = None
    for _ in range(100):
        payload = client.get(f"/api/jobs/{job_id}").json()
        if payload["status"] in {"COMPLETED", "FAILED"}:
            break
        time.sleep(0.05)
    assert payload is not None
    assert payload["status"] == "COMPLETED", payload
    result = client.get(f"/api/jobs/{job_id}/result")
    assert result.status_code == 200
    assert "# Tài liệu demo" in result.json()["markdown"]
    assert result.json()["quality"]["repair_profile_applied"] == "generic"
    bundle = client.get(f"/api/jobs/{job_id}/download/bundle")
    assert bundle.status_code == 200
    assert bundle.content.startswith(b"PK")
    with zipfile.ZipFile(io.BytesIO(bundle.content)) as archive:
        assert set(archive.namelist()) >= {
            "document.normalized.md",
            "quality_report.json",
            "manifest.json",
        }
    update = client.patch(
        f"/api/jobs/{job_id}/markdown",
        json={"content": "# Nội dung đã chỉnh sửa\n"},
    )
    assert update.status_code == 200
    assert client.get(f"/api/jobs/{job_id}/result").json()["markdown"] == (
        "# Nội dung đã chỉnh sửa\n"
    )
    deleted = client.delete(f"/api/jobs/{job_id}")
    assert deleted.status_code == 204
    assert client.get(f"/api/jobs/{job_id}").status_code == 404


def test_invalid_job_id_and_mislabeled_file_are_rejected() -> None:
    assert client.get("/api/jobs/not-a-job").status_code == 404
    response = client.post(
        "/api/jobs",
        data={"mode": "fast"},
        files={"file": ("fake.pdf", b"not a pdf", "application/pdf")},
    )
    assert response.status_code == 400


def test_math_profile_rejects_docx() -> None:
    response = client.post(
        "/api/jobs",
        data={"mode": "auto", "repair_profile": "math_exam"},
        files={
            "file": (
                "math.docx",
                _docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 400
    assert "chỉ áp dụng cho PDF" in response.json()["detail"]
