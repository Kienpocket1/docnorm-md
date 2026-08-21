from __future__ import annotations

import hashlib
import re
import statistics
import unicodedata
from collections import Counter
from collections.abc import Iterable
from dataclasses import dataclass, field
from pathlib import Path

import pymupdf
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from .models import DocumentMetrics, QualityIssue, RepairProfile


class ConversionError(RuntimeError):
    code = "CONVERSION_FAILED"


class UnsupportedFileError(ConversionError):
    code = "UNSUPPORTED_FILE"


class ScanRequiresOcrError(ConversionError):
    code = "SCAN_REQUIRES_OCR"


class PipelineUnavailableError(ConversionError):
    code = "PIPELINE_UNAVAILABLE"


class HybridUnavailableError(ConversionError):
    code = "HYBRID_UNAVAILABLE"


class HybridWorkerError(ConversionError):
    code = "HYBRID_WORKER_FAILED"


class MathVlmWorkerError(ConversionError):
    code = "MATH_VLM_WORKER_FAILED"


class OutputQualityError(ConversionError):
    code = "OUTPUT_QUALITY_FAILED"


@dataclass(slots=True)
class ConversionDocument:
    markdown: str
    metrics: DocumentMetrics
    issues: list[QualityIssue]
    engine_chain: list[str]
    backend: str
    quality_disposition: str
    source_sha256: str
    assets: list[Path] = field(default_factory=list)
    repair_profile_requested: RepairProfile = RepairProfile.AUTO
    repair_profile_applied: RepairProfile = RepairProfile.GENERIC
    repair_version: str | None = None
    repair_fingerprint_score: float = 0.0
    resolved_issue_codes: list[str] = field(default_factory=list)
    page_markdown: dict[int, str] = field(default_factory=dict)


@dataclass(frozen=True, slots=True)
class PdfTextLayerAnalysis:
    page_count: int
    text_pages: int
    image_dominant_pages: int
    total_alphanumeric_characters: int
    requires_ocr: bool

    @property
    def text_page_coverage(self) -> float:
        return self.text_pages / self.page_count if self.page_count else 0.0


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def clean_markdown(value: str) -> str:
    text = unicodedata.normalize("NFC", value).replace("\r\n", "\n")
    text = text.replace("\r", "\n")
    lines: list[str] = []
    double_bullet = re.compile(r"^(\s*)[-+*]\s+[-+*]\s+")
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        line = double_bullet.sub(r"\1- ", line)
        lines.append(line)
    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return text + "\n" if text else ""


def _escape_cell(value: str) -> str:
    return (
        unicodedata.normalize("NFC", value)
        .replace("\r", "")
        .replace("\n", "<br>")
        .replace("|", "\\|")
        .strip()
    )


def _render_docx_table(table: Table) -> str:
    rows = [[_escape_cell(cell.text) for cell in row.cells] for row in table.rows]
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    output = [
        "| " + " | ".join(rows[0]) + " |",
        "| " + " | ".join("---" for _ in range(width)) + " |",
    ]
    output.extend("| " + " | ".join(row) + " |" for row in rows[1:])
    return "\n".join(output)


def _iter_docx_blocks(document: DocxDocument) -> Iterable[Paragraph | Table]:
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _docx_heading_level(paragraph: Paragraph) -> int | None:
    try:
        style = (paragraph.style.name or "").strip()
    except (AttributeError, KeyError):
        style = ""
    match = re.match(r"^(?:heading|tiêu đề)\s*(\d+)$", style, re.IGNORECASE)
    if match:
        return min(6, max(1, int(match.group(1))))
    return None


def _docx_is_list(paragraph: Paragraph) -> tuple[bool, bool]:
    properties = paragraph._p.pPr
    numbering = properties.numPr if properties is not None else None
    try:
        style = (paragraph.style.name or "").casefold()
    except (AttributeError, KeyError):
        style = ""
    is_list = numbering is not None or style.startswith(("list", "danh sách"))
    return is_list, "number" in style or "số" in style


def extract_docx_basic(source: Path) -> ConversionDocument:
    document = Document(str(source))
    sections: list[str] = []
    counts: Counter[str] = Counter()
    list_number = 0
    for block in _iter_docx_blocks(document):
        if isinstance(block, Table):
            rendered = _render_docx_table(block)
            if rendered:
                sections.append(rendered)
                counts["tables"] += 1
            continue
        text = unicodedata.normalize("NFC", block.text).strip()
        if not text:
            continue
        level = _docx_heading_level(block)
        if level is not None:
            sections.append(f"{'#' * level} {text}")
            counts["headings"] += 1
            list_number = 0
            continue
        is_list, numbered = _docx_is_list(block)
        if is_list:
            list_number += 1
            marker = f"{list_number}." if numbered else "-"
            sections.append(f"{marker} {text}")
            counts["lists"] += 1
        else:
            sections.append(text)
            counts["paragraphs"] += 1
            list_number = 0
    markdown = clean_markdown("\n\n".join(sections))
    if not markdown:
        raise ConversionError("DOCX không chứa nội dung có thể trích xuất")
    images = len(document.inline_shapes)
    metrics = DocumentMetrics(
        pages=None,
        elements=sum(counts.values()),
        headings=counts["headings"],
        paragraphs=counts["paragraphs"],
        tables=counts["tables"],
        lists=counts["lists"],
        images=images,
        markdown_characters=len(markdown),
        broken_characters=markdown.count("\ufffd"),
        provenance_coverage=1.0,
    )
    issue = QualityIssue(
        severity="WARNING",
        code="BASIC_FALLBACK_LIMITED_FIDELITY",
        message=(
            "Đang dùng bộ đọc DOCX cơ bản; provenance cấp trang không có trong DOCX."
        ),
    )
    return ConversionDocument(
        markdown=markdown,
        metrics=metrics,
        issues=[issue],
        engine_chain=["DOCX_BASIC"],
        backend="BASIC_FALLBACK",
        quality_disposition="PASS_WITH_WARNINGS",
        source_sha256=sha256_file(source),
    )


def _pdf_text_blocks(page: pymupdf.Page) -> list[tuple[str, float, bool]]:
    output: list[tuple[str, float, bool]] = []
    payload = page.get_text("dict", sort=True)
    for block in payload.get("blocks", []):
        if block.get("type") != 0:
            continue
        lines: list[str] = []
        sizes: list[float] = []
        bold = False
        for line in block.get("lines", []):
            parts: list[str] = []
            for span in line.get("spans", []):
                text = str(span.get("text") or "").strip()
                if text:
                    parts.append(text)
                    sizes.append(float(span.get("size") or 0))
                    bold = bold or bool(int(span.get("flags") or 0) & 16)
            if parts:
                lines.append(" ".join(parts))
        text = "\n".join(lines).strip()
        if text:
            output.append((text, max(sizes, default=0.0), bold))
    return output


def analyze_pdf_text_layer(source: Path) -> PdfTextLayerAnalysis:
    """Classify image-scanned PDFs without attempting OCR.

    A page is considered text-bearing only when it has enough alphanumeric
    characters to be useful for normalization. A low-text page is considered
    image-dominant when one raster image covers most of the page. The document
    is routed to OCR only for a single-page scan, an all-image scan, or when at
    least 20% of a multi-page document consists of image-dominant scan pages.
    This avoids sending ordinary PDFs with an image-only cover through OCR.
    """

    try:
        document = pymupdf.open(source)
    except Exception as error:
        raise ConversionError(f"Không mở được PDF: {error}") from error
    try:
        if document.page_count < 1:
            raise ConversionError("PDF không có trang")
        text_pages = 0
        image_dominant_pages = 0
        total_alphanumeric = 0
        for page in document:
            text = page.get_text("text", sort=True)
            alphanumeric = sum(character.isalnum() for character in text)
            total_alphanumeric += alphanumeric
            if alphanumeric >= 40:
                text_pages += 1
                continue
            page_area = max(float(page.rect.width * page.rect.height), 1.0)
            dominant_image = False
            try:
                for image in page.get_image_info():
                    bbox = image.get("bbox")
                    if not bbox or len(bbox) != 4:
                        continue
                    image_area = max(
                        0.0,
                        float(bbox[2] - bbox[0]) * float(bbox[3] - bbox[1]),
                    )
                    if image_area / page_area >= 0.60:
                        dominant_image = True
                        break
            except Exception:
                dominant_image = bool(page.get_images(full=True))
            if dominant_image:
                image_dominant_pages += 1
        page_count = document.page_count
    finally:
        document.close()

    scan_ratio = image_dominant_pages / page_count
    requires_ocr = (
        image_dominant_pages == page_count
        or (page_count == 1 and image_dominant_pages == 1)
        or (image_dominant_pages >= 2 and scan_ratio >= 0.20)
    )
    return PdfTextLayerAnalysis(
        page_count=page_count,
        text_pages=text_pages,
        image_dominant_pages=image_dominant_pages,
        total_alphanumeric_characters=total_alphanumeric,
        requires_ocr=requires_ocr,
    )


def extract_pdf_basic(source: Path) -> ConversionDocument:
    try:
        document = pymupdf.open(source)
    except Exception as error:
        raise ConversionError(f"Không mở được PDF: {error}") from error
    if document.page_count < 1:
        raise ConversionError("PDF không có trang")
    page_blocks = [_pdf_text_blocks(page) for page in document]
    all_sizes = [size for blocks in page_blocks for _, size, _ in blocks if size > 0]
    body_size = statistics.median(all_sizes) if all_sizes else 11.0
    page_lines = [
        [
            line.strip()
            for text, _, _ in blocks
            for line in text.splitlines()
            if line.strip()
        ]
        for blocks in page_blocks
    ]
    repeated = Counter(line.casefold() for lines in page_lines for line in set(lines))
    repetition_threshold = max(3, int(document.page_count * 0.6 + 0.5))
    boilerplate = {
        line
        for line, count in repeated.items()
        if count >= repetition_threshold and len(line) <= 120
    }
    sections: list[str] = []
    text_pages = 0
    filtered_boilerplate = 0
    headings = 0
    paragraphs = 0
    images = 0
    for page_number, (page, blocks) in enumerate(zip(document, page_blocks), start=1):
        images += len(page.get_images(full=True))
        page_has_text = False
        page_sections: list[str] = []
        for text, size, bold in blocks:
            normalized_lines = []
            for line in text.splitlines():
                cleaned = " ".join(line.split()).strip()
                if cleaned.casefold() in boilerplate:
                    filtered_boilerplate += 1
                    continue
                if cleaned:
                    normalized_lines.append(cleaned)
            text = " ".join(normalized_lines).strip()
            if not text:
                continue
            page_has_text = True
            if len(text) <= 140 and (size >= body_size + 1.8 or bold):
                level = 1 if page_number == 1 and not sections else 2
                page_sections.append(f"{'#' * level} {text}")
                headings += 1
            else:
                page_sections.append(text)
                paragraphs += 1
        if page_has_text:
            text_pages += 1
            sections.append(f"<!-- page: {page_number} -->")
            sections.extend(page_sections)
    document.close()
    markdown = clean_markdown("\n\n".join(sections))
    if not markdown or sum(ch.isalnum() for ch in markdown) < 20:
        raise ScanRequiresOcrError(
            "PDF gần như không có text layer; cần bật OCR/Hybrid trước khi chuẩn hóa."
        )
    issues = [
        QualityIssue(
            severity="WARNING",
            code="BASIC_FALLBACK_LIMITED_FIDELITY",
            message="Đang dùng PyMuPDF; bảng và cấu trúc phức tạp có thể cần kiểm tra thủ công.",
        )
    ]
    if text_pages < len(page_blocks):
        issues.append(
            QualityIssue(
                severity="WARNING",
                code="PAGES_WITHOUT_TEXT",
                message=f"Có {len(page_blocks) - text_pages} trang không có text layer.",
            )
        )
    if filtered_boilerplate:
        issues.append(
            QualityIssue(
                severity="INFO",
                code="REPEATED_HEADER_FOOTER_FILTERED",
                message=f"Đã lọc {filtered_boilerplate} dòng header/footer lặp.",
            )
        )
    metrics = DocumentMetrics(
        pages=len(page_blocks),
        elements=headings + paragraphs,
        headings=headings,
        paragraphs=paragraphs,
        images=images,
        markdown_characters=len(markdown),
        broken_characters=markdown.count("\ufffd"),
        provenance_coverage=text_pages / len(page_blocks),
    )
    return ConversionDocument(
        markdown=markdown,
        metrics=metrics,
        issues=issues,
        engine_chain=["PYMUPDF_BASIC"],
        backend="BASIC_FALLBACK",
        quality_disposition="PASS_WITH_WARNINGS",
        source_sha256=sha256_file(source),
    )


def extract_basic(source: Path) -> ConversionDocument:
    suffix = source.suffix.casefold()
    if suffix == ".pdf":
        return extract_pdf_basic(source)
    if suffix == ".docx":
        return extract_docx_basic(source)
    raise UnsupportedFileError("Chỉ hỗ trợ PDF và DOCX")
